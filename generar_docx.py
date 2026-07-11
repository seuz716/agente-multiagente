#!/usr/bin/env python3
"""Generador DOCX con formato APA 7 - estudiantil con datos ficticios"""
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def crear_docx_aportada(titulo="Análisis del Sistema Agente-Multiagente", autor="César Alarcón", afiliacion="Universidad Nacional de Ingeniería", curso="Sistemas Operativos I", docente="Ing. Carlos Mendoza", fecha="Julio 2026"):
    doc = Document()
    
    # Márgenes 2.54 cm
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # Portada
    for _ in range(5):
        doc.add_paragraph()
    
    titulo_para = doc.add_paragraph()
    titulo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo_para.add_run(titulo.upper())
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Autor
    autor_para = doc.add_paragraph()
    autor_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = autor_para.add_run(f"\n\n{autor}")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    
    # Afiliación
    afil_para = doc.add_paragraph()
    afil_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = afil_para.add_run(afiliacion)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    
    # Curso
    curso_para = doc.add_paragraph()
    curso_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = curso_para.add_run(f"\n\nCurso: {curso}")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    
    docente_para = doc.add_paragraph()
    docente_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = docente_para.add_run(f"Docente: {docente}")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    
    fecha_para = doc.add_paragraph()
    fecha_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fecha_para.add_run(f"\n\n{fecha}")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    
    # Header con número de página
    section = doc.sections[0]
    header = section.header
    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = para.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    fld.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld)
    
    # Contenido
    doc.add_paragraph()
    sec1 = doc.add_paragraph()
    sec1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = sec1.add_run("1. Introducción")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    
    para = doc.add_paragraph("Este documento presenta el análisis del sistema agente-multiagente desarrollado en Node.js y Django. El sistema implementa una arquitectura de agentes especializados con capacidades de procesamiento de órdenes en lenguaje natural.")
    para.paragraph_format.first_line_indent = Cm(1.27)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    doc.add_paragraph()
    sec2 = doc.add_paragraph()
    sec2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = sec2.add_run("2. Desarrollo")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    
    para = doc.add_paragraph("El sistema consta de tres agentes principales: código, archivos y sistema. Cada agente procesa órdenes específicas mediante análisis de palabras clave y scoring.")
    para.paragraph_format.first_line_indent = Cm(1.27)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    para = doc.add_paragraph("Durante la auditoría se identificaron problemas críticos en el orquestador y vulnerabilidades de seguridad en el servidor.")
    para.paragraph_format.first_line_indent = Cm(1.27)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    doc.add_paragraph()
    sec3 = doc.add_paragraph()
    sec3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = sec3.add_run("3. Conclusiones")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    
    para = doc.add_paragraph("El sistema requiere mejoras en seguridad y autenticación antes de producción.")
    para.paragraph_format.first_line_indent = Cm(1.27)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    doc.save('/home/cesar/Documentos/agente-multiagente/informe_apa.docx')
    return '/home/cesar/Documentos/agente-multiagente/informe_apa.docx'

if __name__ == '__main__':
    # Parse cli args: titulo, autor, afiliacion, curso, docente, fecha
    args = sys.argv[1:]
    # Solo pasar valores si existen
    kwargs = {}
    if len(args) > 0: kwargs['titulo'] = args[0]
    if len(args) > 1: kwargs['autor'] = args[1]
    if len(args) > 2: kwargs['afiliacion'] = args[2]
    if len(args) > 3: kwargs['curso'] = args[3]
    if len(args) > 4: kwargs['docente'] = args[4]
    if len(args) > 5: kwargs['fecha'] = args[5]
    
    ruta = crear_docx_aportada(**kwargs)
    print(f"Creado: {ruta}")